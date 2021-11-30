import string
from docx import Document
from docx.shared import Pt


if __name__ == '__main__':
    font_names = ['Helvetica', 'Calibri', 'Futura', 'Garamond', 'Times New Roman',
                  'Arial', 'Cambria', 'Verdana', 'Rockwell', 'Franklin Gothic']

    for character in string.printable[:94:]:
        doc = Document()
        for font_name in font_names:
            for i in range(8, 19):
                p = doc.add_paragraph('')
                run_n = p.add_run(character)
                run_n.font.size = Pt(i)
                run_n.font.name = font_name
                doc.add_page_break()
                p = doc.add_paragraph('')
                run_b = p.add_run(character)
                run_b.bold = True
                run_b.font.size = Pt(i)
                run_b.font.name = font_name
                doc.add_page_break()
                p = doc.add_paragraph('')
                run_i = p.add_run(character)
                run_i.italic = True
                run_i.font.size = Pt(i)
                run_i.font.name = font_name
                doc.add_page_break()
                p = doc.add_paragraph('')
                run_bi = p.add_run(character)
                run_bi.bold = True
                run_bi.italic = True
                run_bi.font.size = Pt(i)
                run_bi.font.name = font_name
                doc.add_page_break()
        doc.save(f'docx/{character}.docx')
